# -*- coding: utf-8 -*-
import sys

from docx import Document
# import pandas as pd
import re
import sqlite3
from PyQt5.QtSql import QSqlDatabase, QSqlQuery
from PyQt5.QtWidgets import QApplication, QMessageBox, QLabel
import sys
import GUI

doc = Document(r'C:\Users\asus\Documents\NDT YKR\NDT\REPORTS 2021\04-YKR-ON-UTT-21-788 (KUT 560) 8.12ap.docx')

# Извлекаем необходимые данные из  репорта
all_tables = doc.tables
# создаем пустой словарь под данные таблиц
data_tables = {i: None for i in range(len(all_tables))}

for i, table in enumerate(all_tables):
    # создаем список строк для таблицы `i` (пока пустые)
    data_tables[i] = [[] for _ in range(len(table.rows))]
    # проходимся по строкам таблицы `i`
    for j, row in enumerate(table.rows):
        for cell in row.cells:
            data_tables[i][j].append(cell.text)

# Словарь очищенных таблиц с данными
clear_tables = {}
# Число таблиц с очищенными данными
number_clear_tables = 0
# проходим по всем найденным таблицам (спискам списков)

for i in data_tables:
    # проходим по всем  спискам
    for ii in data_tables[i]:
        # проходим по всем элементам списка
        for iii in ii:
            # выбираем только таблицы с нужными данными
            if re.match(r'Nominal thickness', iii) or re.match(r'DIA', iii):
                clear_tables[number_clear_tables] = data_tables[i]
                number_clear_tables += 1

# Извлекаем номер и дату репорта из верхнего колонтитула
h = doc.sections[0].header.tables
# создаем пустой словарь под данные верхнего колонтитула
data_header = {i: None for i in range(len(h))}

for i, table in enumerate(h):
    # создаем список строк для таблицы `i` (пока пустые)
    data_header[i] = [[] for _ in range(len(table.rows))]
    # проходимся по строкам таблицы `i`
    for j, row in enumerate(table.rows):
        for cell in row.cells:
            data_header[i][j].append(cell.text)

pp = 0
# словарь номера репорта, даты репорта, номер ворк ордера
rep_number = {}

for i in data_header[0]:
    p = 0
    for ii in i:
        if re.match(r'Date', ii):
            rep_number['report_date'] = data_header[0][pp][p+1]
        if re.match(r'Report', ii):
            rep_number['report_number'] = data_header[0][pp][p+1]
        if re.match(r'Work', ii):
            rep_number['work_order'] = data_header[0][pp][p+1]
        p += 1
    pp += 1

# количество очищенных таблиц с данными
len_tables = len(clear_tables)

# очистка таблиц
for i in range(len(clear_tables)):
    for ii in clear_tables[i]:
        # очистка от общей первой строки
        if len(set(clear_tables[i][0])) == 1:
            clear_tables[i].pop(0)


# подключаем базу данных
con = QSqlDatabase.addDatabase('QSQLITE')
con.setDatabaseName('Reports_DB')
app = QApplication(sys.argv)
if not con.open():
    QMessageBox(None, 'App Name - Error!', 'Database Error: %s' % con.lastError().databaseText())
    sys(exit(1))
createTableQuery = QSqlQuery()
createTableQuery.exec(
    """
    CREATE TABLE contacts (
        id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE NOT NULL,
        name VARCHAR(40) NOT NULL,
        job VARCHAR(50),
        email VARCHAR(40) NOT NULL
    )
    """
)
