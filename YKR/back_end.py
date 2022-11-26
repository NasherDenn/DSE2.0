# -*- coding: utf-8 -*-

from docx import Document
import openpyxl
from openpyxl.styles import Font, Alignment
import datetime
import re
import sqlite3
from YKR.props import DB_NAME
import os
import logging
import traceback

# получаем имя машины с которой был осуществлён вход в программу
uname = os.environ.get('USERNAME')
# инициализируем logger
logger = logging.getLogger()
logger_with_user = logging.LoggerAdapter(logger, {'user': uname})


# функция для извлечения данных из репортов и записи в базу данных при нажатии на кнопку "Добавить"
def add_table(name_dir):
    # переменная список для дальнейшего преобразования списка списков в список строк выбранных для загрузки файлов docx
    name_dir_docx = []
    for i in name_dir[:-1]:
        name_dir_docx.append(i)
    # присваиваем переменной список найденных файлов с расширением docx
    list_find_docx = name_dir_docx[0]
    # счётчик репортов с нарушением структуры
    distract_structure = 0
    # список репортов с нарушением структуры
    list_distract_structure = []
    # список таблиц, которые не должны быть записаны
    dont_save_tables = []
    # список количества всего таблиц в файле
    all_table = []
    # список загруженного количества таблиц в репорте
    load_table = []
    # список загружаемых репортов
    list_load_report = []
    # список загруженных таблиц в репорте для отчёта по загрузке
    li_rep = []
    # список НЕ записанных таблиц в репорте
    # li_dont_rep = []
    # проверяем найденные репорты
    try:
        for unit_list_find_docx in list_find_docx:
            # работаем только если файл имеет допустимое имя, начинающееся с "04-YKR..."
            if re.findall(r'04-YKR', unit_list_find_docx):
                doc = Document(unit_list_find_docx)
                # Извлекаем номер и дату репорта из верхнего колонтитула
                head_paragraph = doc.sections[0].header.tables
                # создаем пустой словарь под данные верхнего колонтитула
                data_header = {i: None for i in range(0, len(head_paragraph))}
                for i, table in enumerate(head_paragraph):
                    # создаем список строк для таблицы `i` (пока пустые)
                    data_header[i] = [[] for _ in range(0, len(table.rows))]
                    # проходимся по строкам таблицы `i`
                    for j, row in enumerate(table.rows):
                        for cell in row.cells:
                            data_header[i][j].append(cell.text)
                pp = 0
                # словарь номера репорта, даты репорта, номера Work Order
                rep_number = {}
                for i in data_header[0]:
                    p = 0
                    for ii in i:
                        # если есть слово "Date" и любая цифра, то дата находится в этой же ячейке
                        if re.match(r'Date', ii) and re.findall(r'\d', ii):
                            rep_number['report_date'] = data_header[0][pp][p][-11:]
                        elif re.match(r'Date', ii) and re.findall(r'\D', ii):
                            rep_number['report_date'] = data_header[0][pp][p + 1]
                        # если есть слово "Report" и любая цифра, то номер репорта находится в этой же ячейке
                        if re.match(r'Report', ii) and re.findall(r'\d', ii):
                            rep_number['report_number'] = data_header[0][pp][p][11:]
                        elif re.match(r'Report', ii) and re.findall(r'\D', ii):
                            rep_number['report_number'] = data_header[0][pp][p + 1]
                        # если есть слово "Work" и любая цифра, то номер Work Order ордера находится в этой же ячейке
                        if re.match(r'Work', ii) and re.findall(r'\d', ii):
                            rep_number['work_order'] = data_header[0][pp][p][-8:]
                        elif re.match(r'Work', ii) and re.findall(r'\D', ii):
                            rep_number['work_order'] = data_header[0][pp][p + 1]
                        p += 1
                    pp += 1

                # меняем '-' на '_' в названиях репортов в rep_number['report_number']
                rep_number['report_number'] = re.sub('-', '_', rep_number['report_number'])
                # добавляем в список номера репортов для статистики
                list_load_report.append(rep_number['report_number'])
                # извлекаем необходимые данные из репорта
                # переменная со всеми таблицами в репорте
                all_tables = doc.tables
                # создаем пустой словарь под данные таблиц
                data_tables = {i: None for i in range(0, len(all_tables))}
                for i, table in enumerate(all_tables):
                    # создаем список строк для таблицы `i` (пока пустые)
                    data_tables[i] = [[] for _ in range(0, len(table.rows))]
                    # проходимся по строкам таблицы `i`
                    for j, row in enumerate(table.rows):
                        try:
                            for cell in row.cells:
                                data_tables[i][j].append(cell.text)
                        except IndexError as index_error:
                            dont_save_tables.append(rep_number['report_number'])
                            logger_with_user.error(str(rep_number['report_number']))
                            logger_with_user.error(index_error)
                            break

                # словарь таблиц с необходимыми данными
                clear_tables = {}
                # счётчик порядкового номера (количество) таблиц с очищенными данными
                number_clear_tables = 0
                # проходим по всем найденным таблицам (спискам списков)
                for i in data_tables:
                    # проходим по всем спискам
                    # счётчик (break_break) прерывания обхода списков, если нашли в ячейке "DIA", "Diameter",
                    # "Nominal thickness" - для случая, если эти слова записаны в две объединённые строки
                    break_break = 0
                    for ii in data_tables[i]:
                        # проходим по всем элементам списка
                        for iii in ii:
                            # выбираем только таблицы с нужными данными в ячейках которых есть ключевые слова
                            # "DIA", "Diameter", "Nominal thickness"
                            if re.match(r'Nom', iii) or re.match(r'nom', iii) \
                                    or re.match(r'DIA', iii) or re.match(r'Dia', iii):
                                clear_tables[number_clear_tables] = data_tables[i]
                                number_clear_tables += 1
                                break_break += 1
                                # если нашли ключевое слово, то прерываем дальнейший обход ячеек таблицы
                                # и переводим к следующей таблице
                                break
                        if break_break == 1:
                            break

                # очищенный (с начала таблицы) словарь с данными
                clear_table_top = {}
                # фильтруем полученные данные из таблиц, выбирая только данные, находящиеся ниже строки
                # "INSPECTION RESULTS", ключевого слова "Results"
                # переменная номера строки (INSPECTION RESULTS) в таблице после которой идут необходимые данные
                del_res_top = 0
                for i in list(clear_tables.keys()):
                    for ii in range(len(clear_tables[i])):
                        for iii in range(len(clear_tables[i][ii])):
                            if re.findall(r'result|RESULT|Result', clear_tables[i][ii][0]):
                                del_res_top = ii + 1
                    for iiii in range(del_res_top, len(clear_tables[i])):
                        clear_table_top[i] = clear_tables[i][del_res_top:]
                # очищенный (с конца таблицы) словарь с данными
                clear_table_bottom = {}
                # фильтруем отфильтрованную таблицу сверху (clear_table_top), выбирая только данные, находящиеся выше
                # строки "Examined by", ключевого слова "Exam"
                for i in list(clear_table_top.keys()):
                    del_res_bottom = len(clear_table_top[i])
                    for ii in range(len(clear_table_top[i])):
                        for iii in range(len(clear_table_top[i][ii])):
                            if re.findall(r'Examined by', clear_table_top[i][ii][iii]):
                                del_res_bottom = ii
                    for iiii in range(del_res_bottom):
                        clear_table_bottom[i] = clear_table_top[i][:del_res_bottom]

                # очищаем полученные таблицы (clear_table_bottom) на данном этапе от пустых строк из-за наличия картинок
                for i in list(clear_table_bottom.keys()):
                    # список номер пустых строк
                    n_sp_str = []
                    # активатор наличия пустых строк
                    ch = 0
                    for ii in range(len(clear_table_bottom[i])):
                        if len(set(clear_table_bottom[i][ii])) == 1 and '' in set(clear_table_bottom[i][ii]):
                            # обновляем список пустых строк
                            n_sp_str.append(ii)
                            # активируем счётчик наличия пустых строк
                            ch += 1
                    if ch > 0:
                        p = 0
                        # если пустых строк больше, чем одна
                        if len(n_sp_str) > 1:
                            for d in n_sp_str:
                                # если первая пустая строка
                                if p == 0:
                                    del clear_table_bottom[i][d]
                                    p += 1
                                else:
                                    # все последующие пустые строки
                                    del clear_table_bottom[i][d - p]
                                    p += 1
                        # если пустая строка только первая
                        elif len(n_sp_str) == 1:
                            del clear_table_bottom[i][n_sp_str[0]]

                # очищаем полученные таблицы (clear_table_bottom) на данном этапе от первой ошибочной таблицы из-за
                # наличия "Nominal Thickness", "Diameter" в самой первой таблице
                del_start_table = 0
                for i in list(clear_table_bottom.keys()):
                    for ii in range(len(clear_table_bottom[i])):
                        for iii in range(len(clear_table_bottom[i][ii])):
                            if re.findall(r'Com', clear_table_bottom[i][ii][iii]) \
                                    or re.findall(r'Proced', clear_table_bottom[i][ii][iii]):
                                del_start_table += 1
                if del_start_table > 0:
                    del clear_table_bottom[0]

                # список репортов с нарушением структуры для итоговых данных
                for i in list(clear_table_bottom.keys()):
                    try:
                        for ii in range(len(clear_table_bottom[i])):
                            # проверка, вносился ли уже репорт с таким номером в список
                            if rep_number['report_number'] not in list_distract_structure:
                                if len(clear_table_bottom[i][ii]) == 0:
                                    list_distract_structure.append(rep_number['report_number'])
                                    distract_structure += 1
                    except KeyError as key_error:
                        dont_save_tables.append(rep_number['report_number'])
                        logger_with_user.error(rep_number['report_number'])
                        logger_with_user.error(key_error)
                        break

                # словарь названия столбцов для каждой таблицы
                name_column = {}
                # список номеров колонок с названиями линий
                number_column_line = []
                # список номеров колонок с названиями чертежей
                number_column_drawing = []
                # Выбираем названия столбцов для каждой таблицы
                for i in clear_table_bottom.keys():
                    name_column[i] = clear_table_bottom[i][0]
                    for ii in range(len(clear_table_bottom[i][0])):
                        if re.findall(r'Line|Tag', clear_table_bottom[i][0][ii]):
                            # номер колонки с названием линии для дальнейшего поиска
                            number_column_line.append(ii)
                        if re.findall(r'Draw|Isometr', clear_table_bottom[i][0][ii]):
                            # номер колонки с названием чертежа для дальнейшего поиска
                            number_column_drawing.append(ii)
                # удаляем из таблиц первые списки (названия столбцов)
                for i in clear_table_bottom.keys():
                    del clear_table_bottom[i][0]

                try:
                    # очищаем названия столбцов и приводим к допустимым названиям
                    for i in name_column.keys():
                        for ii in range(len(name_column[i])):
                            # проверяем наличие в названиях столбцов возможные недопустимые имена
                            if re.findall(r'Tag|TAG|Line|LINE', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Line')
                                # номер колонки с номером линии
                            elif re.findall(r'item|Item|description|Description', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Item_description')
                            elif re.findall(r'Section', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Section')
                            elif re.findall(r'Location', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Location')
                            elif re.findall(r'Remark', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Remark')
                            elif re.findall(r'Size|SIZE', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Size')
                            elif re.findall(r'Vertical', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Vertical')
                            elif re.findall(r'Horizontal', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Horizontal')
                            # если DIA x Nom в названии столбца, то записать Nominal_thickness
                            elif re.findall(r'Diametr|Dia|DIA|dia|DIА|Día', name_column[i][ii]) and re.findall(
                                    r'thicknes|Thicknes|Nom', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Nominal_thickness')
                            elif re.findall(r'Diametr|Dia|DIA|dia|DIА|Día', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Diameter')
                            elif re.findall(r'thicknes|Thicknes', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Nominal_thickness')
                            elif re.findall(r'Extrados', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Extrados')
                            elif re.findall(r'Intrados', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Intrados')
                            elif re.findall(r'South|south', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'South')
                            elif re.findall(r'North|north', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'North')
                            elif re.findall(r'West|west', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'West')
                            elif re.findall(r'East|east', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'East')
                            elif re.findall(r'Top|top', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Top')
                            elif re.findall(r'Bottom|bottom', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Bottom')
                            elif re.findall(r'Shell|shell', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Shell')
                            elif re.findall(r'Plate|plate', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Plate')
                            elif re.findall(r'Drawing|drawing|Isometr', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Drawing')
                            elif re.findall(r'Spot', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Spot')
                            elif re.findall(r'Cente', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Center')
                            elif re.findall(r'Row', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Row')
                            elif re.findall(r'Column', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Column')
                            elif re.findall(r'P&ID', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'P_ID')
                            elif re.findall(r'Right', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Right')
                            elif re.findall(r'Left', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Left')
                            elif re.findall(r'Date|date', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Date')
                            elif re.findall(r'Distance', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Distance')
                            elif re.findall(r'Result', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Result')
                            elif re.findall(r'№|S/NO|S/N|s/n|s/no|NO|no', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'S_N')
                            if re.findall(r'/', name_column[i][ii]):
                                b = re.sub(r'/', '_', name_column[i][ii])
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                            if re.findall(r' ', name_column[i][ii]):
                                b = re.sub(r' ', '_', name_column[i][ii])
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                            if re.findall(r'\.', name_column[i][ii]):
                                b = re.sub(r'\.', '_', name_column[i][ii])
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                            if re.findall(r'\n', name_column[i][ii]):
                                # меняем найденное значение
                                b = re.sub(r'\n', '', name_column[i][ii])
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                            if name_column[i][ii].isnumeric():
                                b = '_' + name_column[i][ii] + '_'
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                            if name_column[i][ii][0].isnumeric():
                                b = '_' + name_column[i][ii] + '_'
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                            if re.findall(r'\d+-\d+', name_column[i][ii]):
                                # меняем найденное значение
                                b = '_' + re.sub('-', '_', name_column[i][ii]) + '_'
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                except Exception:
                    # перехватываем все исключения, которые не предусмотрены в коду выше и записываем их в LogFile
                    logger_with_user.error(str(rep_number['report_number']))
                    logger_with_user.error(traceback.format_exc())

                for i in list(clear_table_bottom.keys()):
                    for ii in range(len(clear_table_bottom[i])):
                        for iii in range(len(clear_table_bottom[i][ii])):
                            # очищаем таблицу данных clear_table_bottom
                            if re.findall(r'"|\'\'|”|’’', clear_table_bottom[i][ii][iii]):
                                b = re.sub(r'"|\'\'|”|’’', '', clear_table_bottom[i][ii][iii])
                                clear_table_bottom[i][ii].remove(clear_table_bottom[i][ii][iii])
                                clear_table_bottom[i][ii].insert(iii, b)
                            if re.findall(r'\n', clear_table_bottom[i][ii][iii]):
                                b = re.sub(r'\n', ' ', clear_table_bottom[i][ii][iii])
                                clear_table_bottom[i][ii].remove(clear_table_bottom[i][ii][iii])
                                clear_table_bottom[i][ii].insert(iii, b)
                # активатор наличия номера чертежа в ячейке с номером линии
                sh_drawing = 0
                for i in list(clear_table_bottom.keys()):
                    for ii in range(len(clear_table_bottom[i])):
                        for iii in range(len(clear_table_bottom[i][ii])):
                            # проверяем на наличие в одной ячейке номер линии и номер чертежа
                            if re.findall(r'[AАBВCСDHНMМ]'
                                          r'\d{1,2}'
                                          r'-{1,2}?\s?'
                                          r'\d{3,4}'
                                          r'-?\s?'
                                          r'\D{2}'
                                          r'-?\s?'
                                          r'\d{3}', clear_table_bottom[i][ii][iii]) and re.findall(r'KE01-.+|TR01-.+',
                                                                                                   clear_table_bottom[i][ii][iii]):
                                # если нашли, находим номер чертежа
                                d = re.findall(r'KE01-.+|TR01-.+', clear_table_bottom[i][ii][iii])[0]
                                # удаляем номер чертежа из ячейки с номером линии
                                clear_table_bottom[i][ii][iii] = clear_table_bottom[i][ii][iii].replace(d, '')
                                # удаляем пустые символы, оставшиеся в ячейке
                                clear_table_bottom[i][ii][iii] = clear_table_bottom[i][ii][iii].replace(' ', '')
                                # удаляем "/" символы, оставшиеся в ячейке
                                clear_table_bottom[i][ii][iii] = clear_table_bottom[i][ii][iii].replace('/', '')
                                # добавляем номер чертежа в столбец
                                clear_table_bottom[i][ii].insert(1, d)
                                # активируем наличие номера чертежа в ячейке с номером линии
                                sh_drawing += 1
                    if sh_drawing != 0:
                        name_column[i].insert(1, 'Drawing')

                # проверяем таблицу на тип "сетка", если первые строки имеют только два значения на всю длину списка
                # set == 2 и значения "Line|Tag", "Diameter", "Nominal thickness", "Item description", то таблица является "сеткой"
                # тогда первое значение из set - название столбца, второе - значение на все строки
                # уникальное название столбца в случае репорта "сетка"
                unique_name_column = []
                # уникальное значение для строки в случае репорта "сетка"
                unique_value_table = []
                # список строк на удаление
                list_for_remove = []
                # список номеров таблиц в которых удалять строки
                list_index_remove = []
                # перебираем строки
                for ii in clear_table_bottom.keys():
                    # переменная номеров строк
                    index_set = 0
                    for iii in clear_table_bottom[ii]:
                        # если в строке всего два уникальных названия, то первое - название столбца,
                        # последнее - значение
                        if len(set(iii)) == 2:
                            # первое - iii[:1] - название для столбца
                            # второе - iii[-1:] - значение для строк
                            unique_name_column.insert(0, iii[:1])
                            unique_value_table.insert(0, iii[-1:])
                            list_for_remove.insert(0, clear_table_bottom[ii][index_set])
                            list_index_remove.insert(0, ii)
                        index_set += 1
                # преобразуем список списков в список строк
                for i in range(len(unique_value_table)):
                    j = unique_value_table[i]
                    unique_value_table.remove(j)
                    unique_value_table.insert(i, j[0])
                # удаляем из таблицы строки, которые должны быть названиями столбцов и их значениями
                e = 0
                for i in list_index_remove:
                    clear_table_bottom[i].remove(list_for_remove[e])
                    e += 1
                # перебираем название столбцов
                for i in name_column.keys():
                    # если всего два уникальных названия
                    if len(set(name_column[i])) == 2:
                        # и первое "Line"
                        if re.findall(r'Line|line|Tag|tag|Contr|contr|Objec|objec', name_column[i][0]):
                            # то добавляем в список уникальных столбцов
                            unique_name_column.insert(0, 'Line')
                            # и значений строк
                            # проверяем значение на номер линии и номер чертежа в одной ячейке (последней)
                            # name_column[i][len(name_column[i]) - 1]
                            if re.findall(r'A1|B0|C2|C3|D1|D6|D7|H1|H2|M1|M2.+KE01|TR01',
                                          name_column[i][len(name_column[i]) - 1]):
                                # переменная для отделения номера линии от номера чертежа
                                rebuild = name_column[i][len(name_column[i]) - 1]
                                # определяем индекс с которого начинается номер чертежа
                                if rebuild.find('KE01'):
                                    index_drawing = rebuild.find('KE01')
                                elif rebuild.find('TR01'):
                                    index_drawing = rebuild.find('TR01')
                                # разъединяем номер линии и номер чертежа
                                temp_line = rebuild[:index_drawing]
                                temp_drawing = rebuild[index_drawing:]
                                # меняем "_" на "-"
                                temp_line = re.sub('_', '-', temp_line)
                                temp_drawing = re.sub('_', '-', temp_drawing)
                                # убираем лишние "-" в начале и в конце
                                while temp_line[0] == '-':
                                    temp_line = temp_line[1:]
                                while temp_drawing[0] == '-':
                                    temp_drawing = temp_drawing[1:]
                                while temp_drawing[-1:] == '-':
                                    temp_drawing = temp_drawing[:-1]
                                while temp_line[-1:] == '-':
                                    temp_line = temp_line[:-1]
                                # добавляем в уникальный список значений строк номер линии
                                unique_value_table.insert(0, temp_line)
                                list_index_remove.insert(0, i)
                                # если нашли и номер линии и номер чертежа, то добавляем в уникальный список названий
                                # столбцов "Drawing"
                                unique_name_column.insert(1, 'Drawing')
                                # добавляем в уникальный список значений строк номер чертежа
                                unique_value_table.insert(1, temp_drawing)
                                list_index_remove.insert(1, i)
                            # иначе добавляем в уникальный список значений строк номер линии
                            else:
                                unique_name_column.insert(0, name_column[i][-1])
                                list_index_remove.insert(0, i)

                    # проверяем есть ли уже добавленный столбец "Line" и его значение в таблицу ранее, когда столбец "Line"
                    # не нашёлся в таблице с данными, но номер линии или чертежа был найден по шаблону в головной таблице
                    if len(set(name_column[i])) == 3:
                        pass

                # меняем в unique_name_column списки на строки
                for i in unique_name_column:
                    if type(i) == list:
                        ind = unique_name_column.index(i)
                        unique_name_column.remove(i)
                        unique_name_column.insert(ind, i[0])
                # дополняем название столбцов новыми названиями из unique_name_column, а после вставляем старые
                if unique_name_column:
                    j = 0
                    name_column = {}
                    for i in set(list_index_remove):
                        name_column[i] = []
                    for i in list_index_remove:
                        name_column[i].append(unique_name_column[j])
                        j += 1
                    for i in name_column.keys():
                        for ii in clear_table_bottom[i][0]:
                            name_column[i].append(ii)
                    for i in clear_table_bottom.keys():
                        # удаляем из таблиц clear_table_bottom старые названия столбцов
                        clear_table_bottom[i].pop(0)
                    # приводим названия столбцов к допустимым
                    for i in name_column.keys():
                        for ii in range(len(name_column[i])):
                            # проверяем наличие в названиях столбцов возможные недопустимые имена
                            if re.findall(r'Tag|TAG|Line|LINE', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Line')
                                # номер колонки с номером линии
                            elif re.findall(r'item|Item|description|Description', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Item_description')
                            elif re.findall(r'Section', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Section')
                            elif re.findall(r'Location', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Location')
                            elif re.findall(r'Remark', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Remark')
                            elif re.findall(r'Size|SIZE', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Size')
                            elif re.findall(r'Vertical', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Vertical')
                            elif re.findall(r'Horizontal', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Horizontal')
                            # если DIA x Nom в названии столбца, то записать Nominal_thickness
                            elif re.findall(r'Diametr|Dia|DIA|dia|DIА|Día', name_column[i][ii]) and re.findall(
                                    r'thicknes|Thicknes|Nom', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Nominal_thickness')
                            elif re.findall(r'Diametr|Dia|DIA|dia|DIА|Día', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Diameter')
                            elif re.findall(r'thicknes|Thicknes', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Nominal_thickness')
                            elif re.findall(r'Extrados', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Extrados')
                            elif re.findall(r'Intrados', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Intrados')
                            elif re.findall(r'South|south', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'South')
                            elif re.findall(r'North|north', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'North')
                            elif re.findall(r'West|west', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'West')
                            elif re.findall(r'East|east', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'East')
                            elif re.findall(r'Top|top', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Top')
                            elif re.findall(r'Bottom|bottom', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Bottom')
                            elif re.findall(r'Shell|shell', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Shell')
                            elif re.findall(r'Plate|plate', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Plate')
                            elif re.findall(r'Drawing|drawing|Isometr', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Drawing')
                            elif re.findall(r'Spot', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Spot')
                            elif re.findall(r'Cente', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Center')
                            elif re.findall(r'Row', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Row')
                            elif re.findall(r'Column', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Column')
                            elif re.findall(r'P&ID', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'P_ID')
                            elif re.findall(r'Right', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Right')
                            elif re.findall(r'Left', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Left')
                            elif re.findall(r'Date|date', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Date')
                            elif re.findall(r'Distance', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Distance')
                            elif re.findall(r'Result', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'Result')
                            elif re.findall(r'№|S/NO|S/N|s/n|s/no|NO|no', name_column[i][ii]):
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, 'S_N')
                            if re.findall(r'/', name_column[i][ii]):
                                b = re.sub(r'/', '_', name_column[i][ii])
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                            if re.findall(r' ', name_column[i][ii]):
                                b = re.sub(r' ', '_', name_column[i][ii])
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                            if re.findall(r'\.', name_column[i][ii]):
                                b = re.sub(r'\.', '_', name_column[i][ii])
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                            if re.findall(r'\n', name_column[i][ii]):
                                # меняем найденное значение
                                b = re.sub(r'\n', '', name_column[i][ii])
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                            if name_column[i][ii].isnumeric():
                                b = '_' + name_column[i][ii] + '_'
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                            if name_column[i][ii][0].isnumeric():
                                b = '_' + name_column[i][ii] + '_'
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)
                            if re.findall(r'\d+-\d+', name_column[i][ii]):
                                # меняем найденное значение
                                b = '_' + re.sub('-', '_', name_column[i][ii]) + '_'
                                # удаляем найденное значение
                                name_column[i].remove(name_column[i][ii])
                                # вставляем на удалённое место допустимое название столбца
                                name_column[i].insert(ii, b)

                    # дополняем таблицу clear_table_bottom новыми значениями "Line", "Drawing", Item_description
                    # и т.д. из unique_value_table счётчики порядковых номеров позиций вставки новых данных
                    # из unique_value_table в clear_table_bottom
                    i_0 = 0
                    i_1 = 0
                    i_2 = 0
                    i_3 = 0
                    i_4 = 0
                    i_5 = 0
                    i_6 = 0
                    i_7 = 0
                    i_8 = 0
                    i_9 = 0
                    i_10 = 0
                    for i in range(len(list_index_remove)):
                        if list_index_remove[i] == 0:
                            for ii in clear_table_bottom[list_index_remove[i]]:
                                ii.insert(i_0, unique_value_table[i])
                            i_0 += 1
                        if list_index_remove[i] == 1:
                            for ii in clear_table_bottom[list_index_remove[i]]:
                                ii.insert(i_1, unique_value_table[i])
                            i_1 += 1
                        if list_index_remove[i] == 2:
                            for ii in clear_table_bottom[list_index_remove[i]]:
                                ii.insert(i_2, unique_value_table[i])
                            i_2 += 1
                        if list_index_remove[i] == 3:
                            for ii in clear_table_bottom[list_index_remove[i]]:
                                ii.insert(i_3, unique_value_table[i])
                            i_3 += 1
                        if list_index_remove[i] == 4:
                            for ii in clear_table_bottom[list_index_remove[i]]:
                                ii.insert(i_4, unique_value_table[i])
                            i_4 += 1
                        if list_index_remove[i] == 5:
                            for ii in clear_table_bottom[list_index_remove[i]]:
                                ii.insert(i_5, unique_value_table[i])
                            i_5 += 1
                        if list_index_remove[i] == 6:
                            for ii in clear_table_bottom[list_index_remove[i]]:
                                ii.insert(i_6, unique_value_table[i])
                            i_6 += 1
                        if list_index_remove[i] == 7:
                            for ii in clear_table_bottom[list_index_remove[i]]:
                                ii.insert(i_7, unique_value_table[i])
                            i_7 += 1
                        if list_index_remove[i] == 8:
                            for ii in clear_table_bottom[list_index_remove[i]]:
                                ii.insert(i_8, unique_value_table[i])
                            i_8 += 1
                        if list_index_remove[i] == 9:
                            for ii in clear_table_bottom[list_index_remove[i]]:
                                ii.insert(i_9, unique_value_table[i])
                            i_9 += 1
                        if list_index_remove[i] == 10:
                            for ii in clear_table_bottom[list_index_remove[i]]:
                                ii.insert(i_10, unique_value_table[i])
                            i_10 += 1

                # активатор отсутствия Line в основной таблице
                sh_line = 0
                # проверяем каждую таблицу на наличие столбца Line, если его нет, то ищем колонку в головной
                # таблице и добавляем её в rep_number
                for i in list(name_column.keys()):
                    if 'Line' not in name_column[i]:
                        # ищем в головной таблице 'Line'
                        sh_line += 1
                        for ii in data_tables:
                            for iii in data_tables[ii]:
                                for iiii in iii:
                                    # ищем совпадение с шаблоном номера линии
                                    if re.match(r'[AАBВCСDHНMМ]'
                                                r'\d{1,2}'
                                                r'-{1,2}?\s?'
                                                r'\d{3,4}'
                                                r'-?\s?'
                                                r'\D{2}'
                                                r'-?\s?'
                                                r'\d{3}', iiii):
                                        # если есть перевод на новую строку, то проверяем каждую строку на наличие
                                        # номера линии
                                        if re.findall(r'\n', iiii):
                                            sp = iiii.split('\n')
                                            for j in sp:
                                                if re.match(r'[AАBВCСDHНMМ]'
                                                            r'\d{1,2}'
                                                            r'-{1,2}?\s?'
                                                            r'\d{3,4}'
                                                            r'-?\s?'
                                                            r'\D{2}'
                                                            r'-?\s?'
                                                            r'\d{3}', j):
                                                    # избавляемся от пробельных символов в начале и в конце строки
                                                    j = j.strip()
                                                    # добавляем не достающий символ '-'
                                                    j = j.replace(' ', '-')
                                                    # удаляем символ дюйма
                                                    if re.findall(r'"|\'\'|”|’’', j):
                                                        b = re.findall(r'"|\'\'|”|’’', j)[0]
                                                        j = j.replace(b, '')
                                                    # найденный номер линии в головной таблице
                                                    line_for_head = j
                                        else:
                                            # избавляемся от пробельных символов в начале и в конце строки
                                            iiii = iiii.strip()
                                            # добавляем не достающий символ '-'
                                            iiii = iiii.replace(' ', '-')
                                            # удаляем символ дюйма
                                            if re.findall(r'"|\'\'|”|’’', iiii):
                                                b = re.findall(r'"|\'\'|”|’’', iiii)[0]
                                                iiii = iiii.replace(b, '')
                                            # найденный номер линии в головной таблице
                                            line_for_head = iiii
                        name_column[i].insert(0, 'Line')

                # не работает
                # if sh_line != 0:
                #     for i in list(clear_table_bottom.keys()):
                #         for ii in range(len(clear_table_bottom[i])):
                #             print(clear_table_bottom)
                #             print(clear_table_bottom[i])
                #             print(line_for_head)
                #
                #             # добавляем номер линии в столбец
                #             # clear_table_bottom[i][ii].insert(0, line_for_head)

                # меняем все "," на ".", убираем все (") в clear_table_bottom, для поиска минимального значения в выводимой
                # таблице
                for i in list(clear_table_bottom.keys()):
                    for ii in range(len(clear_table_bottom[i])):
                        for iii in range(len(clear_table_bottom[i][ii])):
                            clear_table_bottom[i][ii][iii] = re.sub(',', '.', clear_table_bottom[i][ii][iii])
                            clear_table_bottom[i][ii][iii] = re.sub('\'+|\'+|”|″|″', '', clear_table_bottom[i][ii][iii])

                # создаём подключение к базе данных
                conn = sqlite3.connect(DB_NAME)
                cur = conn.cursor()
                # переменная количества добавленных репортов
                check_amount_reports = 0
                # список добавленных таблиц для master
                list_add_table = []
                all_table.append(len(clear_table_bottom.keys()))
                # добавляем данные из репорта в базу данных
                for i in list(clear_table_bottom.keys()):
                    # собираем название таблицы
                    name_clear_table = '\'' + '_' + str(i) + '_' + rep_number['report_number'] + '\''
                    # проверяем, есть такая таблица в базе данных, что бы вносимые данные не повторялись
                    if not cur.execute('SELECT * FROM sqlite_master WHERE  name="{}"'.format(name_clear_table[1:-1])).fetchone():
                        try:
                            # увеличиваем количество добавленных таблиц
                            check_amount_reports += 1
                            # добавляем название таблицы в список для master
                            list_add_table.append(name_clear_table)
                            # создаем таблицу с именем name_clear_table и со столбцами name_column[i]
                            cur.execute(
                                'CREATE TABLE IF NOT EXISTS ' + name_clear_table + ' ({})'.format(','.join(name_column[i])))
                            conn.commit()
                        except (sqlite3.OperationalError, KeyError):
                            # уменьшаем количество добавленных таблиц, если что-то пошло не так
                            check_amount_reports -= 1
                            # удаляем название таблицы из списка для master
                            list_add_table.remove(name_clear_table)
                            dont_save_tables.append(name_clear_table)
                            # сохраняем внесённые изменения, если не было ошибок в репорте
                            conn.commit()
                            # перехватываем все исключения, которые не предусмотрены в коду выше и записываем их в LogFile
                            logger_with_user.error(str(rep_number['report_number']))
                            logger_with_user.error(traceback.format_exc())

                        for ii in clear_table_bottom[i]:
                            try:
                                cur.execute(
                                    'INSERT INTO ' + name_clear_table + ' VALUES (%s)' % ','.join('?' * len(ii)), ii)
                                conn.commit()
                            except sqlite3.OperationalError:
                                dont_save_tables.append(name_clear_table)
                                # сохраняем внесённые изменения, если не было ошибок в репорте
                                conn.commit()
                                # перехватываем все исключения, которые не предусмотрены в коду выше и записываем их в LogFile
                                logger_with_user.error(str(rep_number['report_number']))
                                logger_with_user.error(traceback.format_exc())
                        if name_clear_table in dont_save_tables:
                            try:
                                cur.execute('DROP TABLE ' + name_clear_table)
                                conn.commit()
                            except sqlite3.OperationalError:
                                # перехватываем все исключения, которые не предусмотрены в коду выше и записываем их в LogFile
                                logger_with_user.error(str(rep_number['report_number']))
                                logger_with_user.error(traceback.format_exc())
                                continue
                        conn.commit()
                # преобразуем list_add_table из списка в строку для записи в столбец (list_table_report) в master
                str_add_table = ''
                if len(list_add_table) == 1:
                    str_add_table = list_add_table[0]
                else:
                    for j in list_add_table:
                        str_add_table = str_add_table + j + '\n'
                # создаём таблицу master со столбцами из rep_number
                cur.execute(
                    'CREATE TABLE IF NOT EXISTS master (report_number, report_date, work_order, one_of, list_table_report)')
                # активатор наличия репорта в таблице master
                check_report_number = 0
                # перебираем номера репортов, которые есть в таблице master
                for j in cur.execute('SELECT report_number FROM master').fetchall():
                    # если такой репорт есть (сравниваем последний 6 символов репорта - они уникальны)
                    if rep_number['report_number'][-6:] == j[0][-6:]:
                        # то меняем статус активатора
                        check_report_number += 1
                # если репорт записан, то добавляем его в список для подсчёта количества загруженных таблиц для репорта
                if check_amount_reports:
                    load_table.append(check_amount_reports)
                # иначе записываем "0"
                else:
                    load_table.append('0')
                # список для отчёта по загрузке по загруженным таблицам
                if list_add_table:
                    li_rep.append(list_add_table)
                else:
                    li_rep.append(0)
                # если статус активатора НЕ изменён (такой репорт еще не занесён в базу данных)
                if check_report_number == 0:
                    # вносим данные из rep_number в таблицу master
                    if not check_amount_reports == 0:
                        cur.execute(
                            'INSERT INTO master VALUES (?, ?, ?, ?, ?)', (rep_number['report_number'],
                                                                          rep_number['report_date'],
                                                                          rep_number['work_order'],
                                                                          str(check_amount_reports) + '/' + str(
                                                                              len(clear_table_bottom.keys())),
                                                                          str_add_table))
                        conn.commit()

                # если статус активатора изменён (такой репорт уже есть в базе данных)
                if check_report_number > 0:
                    # то находим этот репорт и получаем данные из столбца 'one_of' и 'list_table_report'
                    # получаем номера всех записанных таблиц в виде строки
                    variable_for_add_table_for_master = cur.execute(
                        'SELECT list_table_report FROM master WHERE report_number = "{}"'.format(
                            rep_number['report_number'])).fetchall()[0][0]
                    # получаем данные из колонки 'one_of'
                    one_of_for_plus = cur.execute('SELECT one_of FROM master WHERE report_number = "{}"'.format(
                        rep_number['report_number'])).fetchall()[0][0]
                    # добавляем в список таблиц новую
                    variable_for_add_table_for_master_new = variable_for_add_table_for_master + str_add_table
                    # получаем индекс '/' в строке one_of
                    index_of_slash = one_of_for_plus.index('/')
                    # получаем первое число перед знаком '/' и преобразуем его в число
                    number_load_report_for_plus = int(one_of_for_plus[:index_of_slash])
                    # увеличиваем количество загруженных репортов на один и преобразуем обратно в строку
                    number_load_report_for_plus_new = str(number_load_report_for_plus + check_amount_reports)
                    # формируем новое значение 'one_of' для обновления
                    one_of_for_plus_new = number_load_report_for_plus_new + one_of_for_plus[index_of_slash:]
                    # обновляем ячейку с номерами вновь добавленных таблиц 'list_table_report'
                    cur.execute('UPDATE master SET list_table_report = "{}" WHERE report_number = "{}"'.format(
                        variable_for_add_table_for_master_new, rep_number['report_number']))
                    # обновляем ячейку с количеством загруженных таблиц 'one_of'
                    cur.execute('UPDATE master SET one_of = "{}" WHERE report_number = "{}"'.format(
                        one_of_for_plus_new, rep_number['report_number']))
                    conn.commit()
                logger_with_user.info('Добавление новых репортов в базу данных: ' + rep_number['report_number'])
                # закрываем соединение с базой данной
                conn.close()
        loading_report(len(list_find_docx), list_load_report, load_table, all_table, li_rep)
    except Exception:
        # перехватываем все исключения, которые не предусмотрены в коду выше и записываем их в LogFile
        logger_with_user.error(str(rep_number['report_number']))
        logger_with_user.error(traceback.format_exc())


# len_list_find_docx = len(list_find_docx) = количество загружаемых файлов xlsx
# l_l_r = list_load_report = список загружаемых файлов
# l_t = load_table = список загруженного количества таблиц в репорте
# a_t = all_table = список количество таблиц в файле (репорте)
# l_r = li_rep = список записанных таблиц в репорте
# функция формирования отчёта по загрузке репортов в формате Excel
def loading_report(len_list_find_docx, l_l_r, l_t, a_t, l_r):
    wb = openpyxl.Workbook()
    # делаем активным первый лист
    sheet_loading_report = wb.active
    # закрепляем первую строку
    sheet_loading_report.freeze_panes = 'A2'

    # настраиваем заголовки листа с отчётом
    # высота всей первой строки
    sheet_loading_report.row_dimensions[1].height = 43
    # ширина столбца "A"
    sheet_loading_report.column_dimensions['A'].width = 13
    # стиль текста, размер и жирное выделение
    style_sheet = Font(name='Calibri', size=11, bold=True)
    sheet_loading_report['A1'].font = style_sheet
    # центрируем положение текста в столбце "A" и делаем автоматический перенос слов
    sheet_loading_report['A1'].alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    # текст в ячейке
    sheet_loading_report['A1'] = 'Количество загружаемых репортов'
    # ширина столбца "B"
    sheet_loading_report.column_dimensions['B'].width = 30
    # стиль текста, размер и жирное выделение
    style_sheet = Font(name='Calibri', size=11, bold=True)
    sheet_loading_report['B1'].font = style_sheet
    # центрируем положение текста в столбце "B"
    sheet_loading_report['B1'].alignment = Alignment(horizontal='center', vertical='center')
    # текст в ячейке
    sheet_loading_report['B1'] = 'Список загружаемых репортов'
    # ширина столбца "C"
    sheet_loading_report.column_dimensions['C'].width = 15
    # стиль текста, размер и жирное выделение
    style_sheet = Font(name='Calibri', size=11, bold=True)
    sheet_loading_report['C1'].font = style_sheet
    # центрируем положение текста в столбце "C" и делаем автоматический перенос слов
    sheet_loading_report['C1'].alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    # текст в ячейке
    sheet_loading_report['C1'] = 'Загружено таблиц  / таблиц в файле'
    # ширина столбца "D"
    sheet_loading_report.column_dimensions['D'].width = 30
    # стиль текста, размер и жирное выделение
    style_sheet = Font(name='Calibri', size=11, bold=True)
    sheet_loading_report['D1'].font = style_sheet
    # центрируем положение текста в столбце "D"
    sheet_loading_report['D1'].alignment = Alignment(horizontal='center', vertical='center')
    # текст в ячейке
    sheet_loading_report['D1'] = 'Список загруженных таблиц'

    # указываем количество загружаемых репортов
    sheet_loading_report['A2'].alignment = Alignment(horizontal='center', vertical='center')
    sheet_loading_report['A2'] = str(len_list_find_docx)
    for i in range(len(l_l_r)):
        # записываем загружаемые репорты в столбец и центрируем их в ячейке
        sheet_loading_report['B' + str(i + 2)].alignment = Alignment(horizontal='center', vertical='center')
        sheet_loading_report['B' + str(i + 2)] = str(l_l_r[i])
        # записываем всего "загружено таблиц / таблиц в файле"
        sheet_loading_report['C' + str(i + 2)].alignment = Alignment(horizontal='center', vertical='center')
        sheet_loading_report['C' + str(i + 2)] = str(l_t[i]) + '/' + str(a_t[i])
        if l_r[i] != 0:
            sheet_loading_report['D' + str(i + 2)].alignment = Alignment(horizontal='center', vertical='center',
                                                                         wrap_text=True)
            # если в количество загруженных таблиц из репорта больше, чем 1, то
            if len(l_r[i]) > 1:
                # преобразуем список строк в одну строку через разделитель - пробел
                s_p = ' '.join(l_r[i])
                s_p = re.sub('\'', ' ', s_p)
                sheet_loading_report['D' + str(i + 2)] = s_p
            else:
                sheet_loading_report['D' + str(i + 2)] = str(l_r[i][0][1:-1])
    # дата и время формирования отчёта по загрузке репортов
    date_time_loading_report = datetime.datetime.now().strftime("%Y-%m-%d %H-%M-%S")
    # имя файла отчёта по загрузке отчётов
    name_loading_report = str(date_time_loading_report) + ' Report loading.xlsx'
    # переименовываем активный лист
    sheet_loading_report.title = str(date_time_loading_report)
    # сохраняем отчёт о загрузке в папку 'Loading report', находящийся в том же каталоге (os.path.abspath(os.getcwd())),
    # что и файл exe
    # отчёты сохраняем по месяцам в году, что бы не накапливались "снежным" комом в одной папке
    # если 'Loading report' не существует,
    new_path = os.path.abspath(os.getcwd()) + '\\Loading report\\' + date_time_loading_report[:7] + '\\'
    if not os.path.exists(new_path):
        # то создаём эту папку
        os.makedirs(new_path)
    wb.save(new_path + name_loading_report)
    # открываем только что сохранённый файл с отчётом о загрузке репортов
    os.startfile(new_path + name_loading_report)
    # закрываем книгу
    wb.close()
    logger_with_user.info('Сформирован файл с отчётом о загрузке репортов\n' + new_path + name_loading_report)


def main():
    add_table()


if __name__ == '__main__':
    main()
